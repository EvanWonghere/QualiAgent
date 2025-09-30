# backend/services.py
import os
import json
from typing import Optional

from dotenv import load_dotenv
from openai import OpenAI
import numpy as np
from sqlalchemy.orm import Session
from docx import Document
import tempfile

from backend import schemas
from backend.models import Chunk, Code, Transcript, Memo

load_dotenv()


# --- Dataset and AI Analysis Services ---
def read_docx_bytes(file_bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
        f.write(file_bytes)
        path = f.name
    doc = Document(path)
    os.remove(path)
    return "\n".join(p.text for p in doc.paragraphs)


def normalize_text(s: str):
    return " ".join(s.strip().split())


def chunk_text(text, approx_tokens: int = None, overlap_ratio=0.1):
    if approx_tokens is None:
        approx_tokens = int(os.getenv("CHUNK_TOKENS", 400))
    avg_char_per_token = 4
    chunk_size = approx_tokens * avg_char_per_token
    overlap = int(chunk_size * overlap_ratio)
    chunks = []
    i = 0
    n = len(text)
    while i < n:
        end = min(n, i + chunk_size)
        chunk = text[i:end]
        chunks.append((i, end, normalize_text(chunk)))
        i = end - overlap
    return chunks


def get_embedding(text: str, config: Optional[schemas.AIConfig] = None):
    request_client = get_openai_client(config)
    # Get model from config, or fallback to environment variable
    model = (config and config.embed_model) or os.getenv("OPENAI_EMBED_MODEL", "text-embedding-3-small")
    resp = request_client.embeddings.create(model=model, input=text)
    return resp.data[0].embedding


def stream_chunks_from_file(path, approx_tokens: int = None, overlap_ratio=0.1):
    if approx_tokens is None:
        approx_tokens = int(os.getenv("CHUNK_TOKENS", 400))
    """
    Reads a large text file from a path and yields its content in smaller,
    overlapping chunks without loading the whole file into memory.
    """
    avg_char_per_token = 4
    chunk_size = approx_tokens * avg_char_per_token
    overlap = int(chunk_size * overlap_ratio)

    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        buffer = ""
        pos = 0
        while True:
            data = f.read(chunk_size)
            if not data:
                # emit remaining buffer if any
                if buffer.strip():
                    start = pos
                    end = pos + len(buffer)
                    yield start, end, normalize_text(buffer)
                break
            buffer += data
            # emit a chunk
            chunk = buffer[:chunk_size]
            start = pos
            end = pos + len(chunk)
            yield start, end, normalize_text(chunk)
            # Prepare buffer for next round with overlap
            buffer = buffer[chunk_size - overlap:]
            pos = end - overlap


def search_similar(db: Session, transcript_id: int, query: str, top_k=5, config: Optional[schemas.AIConfig] = None):
    rows = db.query(Chunk).filter(Chunk.transcript_id == transcript_id).all()
    if not rows: return []
    q_emb = np.array(get_embedding(query, config=config), dtype=float)
    ids, texts, embs = [], [], []
    for r in rows:
        ids.append(r.id)
        texts.append(r.text)
        embs.append(np.array(json.loads(r.embedding), dtype=float))
    if not embs: return []
    embs = np.vstack(embs)
    def cos_sim(a, B):
        a_norm = a / np.linalg.norm(a)
        B_norm = B / np.linalg.norm(B, axis=1, keepdims=True)
        return B_norm @ a_norm
    sims = cos_sim(q_emb, embs)
    top_idx = np.argsort(sims)[-top_k:][::-1]
    results = [{"chunk_id": ids[i], "text": texts[i], "score": float(sims[i])} for i in top_idx]
    # 🧹 CLEANUP: Removed db.close()
    return results


def analyze_chunk_with_llm(chunk_text: str, config: Optional[schemas.AIConfig] = None):
    request_client = get_openai_client(config)
    # Get model from config, or fallback to environment variable
    model = (config and config.llm_model) or os.getenv("OPENAI_LLM_MODEL", "gpt-4o-mini")
    system = "You are a qualitative research assistant. Produce a JSON object with keys: 'summary' (short), 'codes' (list of objects with 'code', 'definition', and 'quotes' list). Output JSON only."
    prompt = f"Transcript chunk:\n\"\"\"{chunk_text}\"\"\"\nPlease produce:\n1) short summary (1-2 sentences)\n2) list up to 5 codes. For each code give: 'code' (short label), 'definition' (one line), and 1-2 short quotes from the chunk that illustrate it.\nReturn JSON only. "
    try:
        res = request_client.chat.completions.create(model=model, messages=[{"role": "system", "content": system},
                                                                        {"role": "user", "content": prompt}],
                                             temperature=0.0, response_format={"type": "json_object"})
        content = res.choices[0].message.content
        return json.loads(content)
    except Exception as e:
        print(f"Error analyzing chunk with LLM: {e}");
        return {"error": str(e)}


# ✨ --- NEW: Robust, recursive formatter now lives in the service layer ---
def format_data_to_markdown(data, indent_level=0) -> str:
    """
    Recursively formats nested dictionaries and lists into a clean Markdown string.
    """
    indent = "  " * indent_level
    if isinstance(data, dict):
        parts = []
        for key, value in data.items():
            key_formatted = key.replace('_', ' ').title()
            prefix = f"{indent}- **{key_formatted}:** " if not isinstance(value, (dict, list)) else f"{indent}**{key_formatted}:**"
            parts.append(f"{prefix}{format_data_to_markdown(value, indent_level + 1)}")
        return "\n".join(parts)
    elif isinstance(data, list):
        return "\n".join(f"{indent}- {format_data_to_markdown(item, indent_level)}" for item in data)
    else:
        return str(data)


def generate_memo_content(db: Session, transcript_id: int, config: Optional[schemas.AIConfig] = None):
    request_client = get_openai_client(config)
    # Get model from config, or fallback to environment variable
    model = (config and config.llm_model) or os.getenv("OPENAI_LLM_MODEL", "gpt-4o-mini")
    chunks = db.query(Chunk).filter(Chunk.transcript_id == transcript_id).limit(15).all()
    if not chunks:
        db.close()
        return {"error": "This transcript has not been processed for AI analysis yet. Chunks are missing."}
    texts = [c.text for c in chunks]
    # db.close()
    if not texts: return {"summary": "No data to generate memo.", "contradictions": [], "followups": []}
    full_text_sample = "\n---\n".join(texts)
    system_prompt = "You are a qualitative research analyst. Your task is to write an analytic memo based on interview excerpts. Your output must be a valid JSON object."
    user_prompt = f"Based on the following excerpts...\n---\n{full_text_sample}\n---\nWrite an analytic memo with three sections... JSON object with the keys 'summary', 'contradictions', and 'followups'..."
    try:
        res = request_client.chat.completions.create(model=model, messages=[{"role": "system", "content": system_prompt},
                                                                        {"role": "user", "content": user_prompt}],
                                             temperature=0.7, response_format={"type": "json_object"})
        content = res.choices[0].message.content
        return json.loads(content)
    except Exception as e:
        print(f"Error generating memo: {e}")
        return {"error": "API call to generate memo failed."}


# ✨ --- NEW HELPER: Shared logic for creating the final memo content string ---
def get_formatted_memo_content(db: Session, transcript_id: int, config: Optional[schemas.AIConfig] = None) -> (str, dict):
    """
    Gets the raw AI response and formats it into a clean Markdown string.
    Returns both the final string and the original JSON for flexibility.
    """
    memo_json = generate_memo_content(db=db, transcript_id=transcript_id, config=config)
    if "error" in memo_json:
        return None, memo_json

    full_content_parts = []
    if "summary" in memo_json:
        full_content_parts.append("## Summary")
        full_content_parts.append(format_data_to_markdown(memo_json["summary"]))
    if "contradictions" in memo_json:
        full_content_parts.append("\n## Contradictions")
        full_content_parts.append(format_data_to_markdown(memo_json["contradictions"]))
    if "followups" in memo_json:
        full_content_parts.append("\n## Follow-up Questions")
        full_content_parts.append(format_data_to_markdown(memo_json["followups"]))

    return "\n".join(full_content_parts), memo_json


# ✅ FINAL FIX: This is the definitive corrected function.
def create_transcript_entry(db: Session, title: str, file_path: str):
    base_title = title
    counter = 1
    while db.query(Transcript).filter(Transcript.title == title).first():
        name, ext = os.path.splitext(base_title)
        title = f"{name}_{counter}{ext}"
        counter += 1

    transcript_db = Transcript(title=title, file_path=file_path)
    db.add(transcript_db)
    db.commit()
    db.refresh(transcript_db)
    return transcript_db


def read_docx_from_path(path):
    # (This function is unchanged)
    doc = Document(path)
    text = "\n".join(p.text for p in doc.paragraphs)
    # Create a new temp file for the text content
    with tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w", encoding="utf-8") as tmp:
        tmp.write(text)
    return tmp.name


def process_transcript_for_ai(db: Session, transcript_id: int):
    """ Memory-safe processing. Reads from the file path stored in the Transcript."""
    transcript = db.query(Transcript).filter(Transcript.id == transcript_id).first()
    if not transcript or not os.path.exists(transcript.file_path):
        raise ValueError("Transcript or its associated file not found")

    # ✨ Update status to show work is in progress
    transcript.status = "processing"
    db.commit()

    try:
        db.query(Chunk).filter(Chunk.transcript_id == transcript_id).delete()

        path_to_process = transcript.file_path
        is_converted_temp = False
        if path_to_process.lower().endswith(".docx"):
            path_to_process = read_docx_from_path(path_to_process)
            is_converted_temp = True

        for start, end, chunk_text_ in stream_chunks_from_file(path_to_process):
            emb = get_embedding(chunk_text_)
            c = Chunk(transcript_id=transcript_id, text=chunk_text_, embedding=json.dumps(emb), start_pos=start,
                      end_pos=end)
            db.add(c)

        if is_converted_temp:
            os.remove(path_to_process)

        transcript.status = "processed"
        db.commit()
        return {"message": f"Transcript '{transcript.title}' processed for AI analysis."}
    except Exception as e:
        # ✨ If anything goes wrong, mark the status as failed
        transcript.status = "failed"
        db.commit()
        # Re-raise the exception to be caught by the endpoint
        raise e
    finally:
        pass

# ✨ --- 新增和修改的函数 ---

def create_memo_from_ai(db: Session, transcript_id: int, config: Optional[schemas.AIConfig] = None):
    """ ✅ FIX: This function now uses the new shared helper to get clean Markdown content."""
    transcript = db.query(Transcript).filter(Transcript.id == transcript_id).first()
    if not transcript:
        return None

    # Use the helper to get the formatted string
    formatted_content, _ = get_formatted_memo_content(db, transcript_id, config=config)
    if not formatted_content:
        return None

    # Save the clean Markdown string to the database
    new_memo = Memo(
        title=f"AI Memo for Transcript: '{transcript.title}'",
        content=formatted_content
    )
    db.add(new_memo)
    db.commit()
    db.refresh(new_memo)
    return new_memo



def generate_and_save_codes(db: Session, transcript_id: int, config: Optional[schemas.AIConfig] = None):
    transcript = db.query(Transcript).filter(Transcript.id == transcript_id).first()
    if not transcript:
        # db.close()
        return {"error": "transcript not found"}

    chunks = db.query(Chunk).filter(Chunk.transcript_id == transcript_id).all()
    saved_codes_count = 0
    for chunk in chunks:
        analysis = analyze_chunk_with_llm(chunk.text, config=config)
        if "codes" in analysis and isinstance(analysis["codes"], list):
            for code_data in analysis["codes"]:
                # AI返回的quotes是一个列表，我们将其合并
                excerpt = "\n".join(code_data.get("quotes", []))
                if not excerpt:  # 如果没有引文，使用部分chunk文本
                    excerpt = chunk.text[:250] + "..."

                new_code = Code(
                    code=code_data.get("code", "Untitled"),
                    excerpt=excerpt,
                    transcript_id=transcript_id  # 关联到Dataset
                )
                db.add(new_code)
                saved_codes_count += 1

    db.commit()
    # db.close()
    return {"message": f"Successfully generated and saved {saved_codes_count} codes for transcript."}


# --- Manual CRUD Services ---


def list_transcripts(db: Session):
    return db.query(Transcript).all()


def create_memo(db: Session, title: str, content: str):
    memo = Memo(title=title, content=content)
    db.add(memo)
    db.commit()
    db.refresh(memo)
    # db.close()
    return memo


def list_memos(db: Session):
    return db.query(Memo).all()


def create_code(db: Session, payload: schemas.CodeCreate):
    new_code = Code(**payload.model_dump())
    # ✨ Corrected validation
    if not new_code.transcript_id and not new_code.memo_id:
        raise ValueError("Code must reference a transcript or a memo")
    db.add(new_code)
    db.commit()
    db.refresh(new_code)
    # db.close()
    return new_code


def list_codes(db: Session):
    codes = db.query(Code).all()
    results = []
    for c in codes:
        source_title = "N/A"
        if c.transcript:
            source_title = f"Transcript: {c.transcript.title}"
        elif c.memo:
            source_title = f"Memo: {c.memo.title}"

        results.append({
            "id": c.id, "code": c.code, "excerpt": c.excerpt, "source": source_title,
            "created_at": c.created_at, "transcript_id": c.transcript_id,
            "memo_id": c.memo_id
        })
    # db.close()
    return results

def delete_transcript(db: Session, transcript_id: int):
    item = db.query(Transcript).filter(Transcript.id == transcript_id).first()
    if item:
        db.delete(item); db.commit()
        return {"deleted": True}
    return {"deleted": False}


def delete_memo(db: Session, memo_id: int):
    item = db.query(Memo).filter(Memo.id == memo_id).first()
    if item:
        db.delete(item); db.commit()
        return {"deleted": True}
    return {"deleted": False}


def delete_code(db: Session, code_id: int):
    code = db.query(Code).filter(Code.id == code_id).first()
    if code:
        db.delete(code)
        db.commit()
        # db.close()
        return {"deleted": True}
    # db.close()
    return {"deleted": False}


# ✨ --- NEW: Functions to get single items by ID ---

def get_transcript_by_id(db: Session, transcript_id: int):
    """Fetches a single transcript and reads its content from the stored file path."""
    transcript = db.query(Transcript).filter(Transcript.id == transcript_id).first()
    if not transcript:
        return None

    content = ""
    try:
        path_to_process = transcript.file_path
        if not os.path.exists(path_to_process):
            raise FileNotFoundError("The source file for this transcript is missing.")

        if path_to_process.lower().endswith(".docx"):
            temp_txt_path = read_docx_from_path(path_to_process)
            with open(temp_txt_path, 'r', encoding='utf-8') as f:
                content = f.read()
            os.remove(temp_txt_path)
        else:
            with open(path_to_process, 'r', encoding='utf-8') as f:
                content = f.read()
    except Exception as e:
        content = f"Error reading file content: {e}"

    # ✅ FIX: Pass the 'status' field when creating the response model.
    return schemas.TranscriptDetail(
        id=transcript.id,
        title=transcript.title,
        status=transcript.status, # This was the missing field
        content=content
    )


def get_memo_by_id(db: Session, memo_id: int):
    """Fetches a single memo by its ID."""
    memo = db.query(Memo).filter(Memo.id == memo_id).first()
    # Use the Pydantic model to prevent session errors
    if memo:
        return schemas.MemoDetail(id=memo.id, title=memo.title, content=memo.content)
    return None

# --- Helper to get a configured OpenAI client ---
def get_openai_client(config: Optional[schemas.AIConfig] = None) -> OpenAI:
    """Creates an OpenAI client based on user-provided config, falling back to .env"""
    # Use user-provided key if available
    api_key = config.api_key if config and config.api_key else os.getenv("OPENAI_API_KEY")
    base_url = config.base_url if config and config.base_url else os.getenv("OPENAI_API_BASE_URL")

    if not api_key:
        raise ValueError("OpenAI API key is not configured.")

    return OpenAI(api_key=api_key, base_url=base_url)


def get_default_config():
    """Reads AI configuration from environment variables."""
    api_key_present = bool(os.getenv("OPENAI_API_KEY"))
    base_url = os.getenv("OPENAI_API_BASE_URL")
    llm_model = os.getenv("OPENAI_LLM_MODEL")
    embed_model = os.getenv("OPENAI_EMBED_MODEL")

    chunk_tokens_str = os.getenv("CHUNK_TOKENS")
    chunk_tokens = int(chunk_tokens_str) if chunk_tokens_str and chunk_tokens_str.isdigit() else None

    return {
        "api_key_set": api_key_present,
        "base_url": base_url,
        "llm_model": llm_model,
        "embed_model": embed_model,
        "chunk_tokens": chunk_tokens
    }
